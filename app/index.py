import json
import math, utils, cloudinary.uploader, admin, hashlib, os, mail
import math, utils, cloudinary.uploader, admin, hashlib, os, dao
from datetime import timedelta

import docx
from models import *
from __init__ import app, loginMNG, db
from flask import render_template, request, redirect, url_for, jsonify, flash, send_file, Flask, session
from docx import Document
from sqlalchemy import func
from flask_login import login_user, logout_user, current_user, login_required
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from sendgrid.helpers.mail import Mail
from sendgrid import SendGridAPIClient
from app.vnpay.form import PaymentForm
from app.vnpay.vnpay import Vnpay

@app.context_processor
def default_response():
    return {
        'category': utils.get_category(),
        'related_books': utils.get_book_home()
    }


@app.route("/")
def index():
    return render_template('home.html')


@app.route("/products")
def list_book():
    cate_id = request.args.get('category_id')
    page = int(request.args.get('page', 1))
    kw = request.args.get('keyword')
    sach = utils.get_book_pagination(kw=kw, category_id=cate_id, page=page)

    return render_template('products.html',
                           page=math.ceil(utils.count_book(category_id=cate_id, kw=kw) / app.config['LIST_SIZE']),
                           lb=sach)


@app.route("/register", methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        flash('Bạn đã đăng ký thành công, vui lòng đăng nhập để tiếp tục mua sắm ^^')
        return redirect(url_for('index'))

    err_msg = ''
    if request.method == 'POST':
        name = request.form.get('registerName')
        email = request.form.get('registerEmail')
        username = request.form.get('registerUserName')
        password = request.form.get('registerPassword')
        confirm = request.form.get('confirmPassword')

        try:
            if password != confirm:
                flash('Mật khẩu không khớp', 'error')

            else:
                if utils.get_user_by_username(username=username):
                    flash('Username đã được sử dung, vui lòng đăng ký username khác', 'error')
                else:
                    if utils.get_user_by_email(email=email):
                        flash('Email đã được đăng ký ở nơi khác, vui lòng đăng ký bằng email khác', 'error')
                    else:
                        # Tạo token xác thực email
                        token = mail.generate_token(email)
                        confirm_url = url_for('confirm_email', token=token, _external=True)
                        flash('Vui lòng kiểm tra email để kích hoạt tài khoản.', 'info')

                        # Gửi email qua SendGrid
                        message = Mail(
                            from_email='bookstore2k4@gmail.com',
                            to_emails=email,
                            subject='Xác nhận email đăng ký',
                            html_content=f'<p>Nhấn vào đường dẫn sau để kích hoạt tài khoản:</p> <a href="{confirm_url}">{confirm_url}</a>'
                        )

                        sg = SendGridAPIClient(app.config["API_KEY"])
                        response = sg.send(message)
                        if response.status_code == 202:
                            utils.add_user(name=name, username=username,
                                           password=password,
                                           email=email,
                                           avatar=url_for('static', filename='images/user.png'))
                            return redirect(url_for('login'))
                        else:
                            flash('Không thể gửi email xác nhận. Vui lòng thử lại sau.')

        except Exception as ex:
            flash('Hệ thống đang có lỗi: ' + str(ex))

    return render_template('register.html', err_msg=err_msg)


@app.route('/confirm_email/<token>', methods=['GET'])
def confirm_email(token):
    try:
        email = mail.confirm_token(token)
        # Lưu người dùng vào DB sau khi email được xác minh
        user = utils.get_user_by_email(email=email)
        user.active = True
        db.session.add(user)
        db.session.commit()
        flash('Tài khoản đã được xác minh và đăng ký thành công!', 'success')
        return redirect(url_for('login'))
    except Exception as e:
        flash('Token xác minh không hợp lệ hoặc hết hạn.', 'error')
        return redirect(url_for('register'))


@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    err_msg = ''
    if request.method.__eq__('POST'):
        username = request.form.get('username')
        password = request.form.get('password')
        if utils.get_user_by_username(username=username):

            user = utils.get_user_by_username(username=username)
            if str(hashlib.md5(password.strip().encode('utf-8')).hexdigest()).__eq__(user.password):
                email = user.email
                if user.active == False:
                    token = mail.generate_token(email)
                    confirm_url = url_for('confirm_email', token=token, _external=True)
                    flash('Vui lòng kiểm tra email để kích hoạt tài khoản.', 'info')

                    # Gửi email qua SendGrid
                    message = Mail(
                        from_email='bookstore2k4@gmail.com',  # Thay bằng email đã xác thực trên SendGrid
                        to_emails=email,
                        subject='Xác nhận email đăng ký',
                        html_content=f'<p>Nhấn vào đường dẫn sau để kích hoạt tài khoản:</p> <a href="{confirm_url}">{confirm_url}</a>'
                    )

                    sg = SendGridAPIClient(app.config["API_KEY"])
                    response = sg.send(message)
                    if response.status_code == 202:
                        return redirect(url_for('login'))
                    else:
                        flash('Không thể gửi email xác nhận. Vui lòng thử lại sau.', 'error')
                else:
                    login_user(user=user)
                    return redirect(url_for('index'))
            else:
                flash('Mật khẩu không chính xác, vui lòng nhập lại.', 'error')
        else:
            flash('Username không tồn tài hoặc không chính xác, vui lòng nhập lại.', 'error')
    return render_template('login.html')


@app.route('/searchAccount', methods=['GET', 'POST'])
def search_account():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    if request.method.__eq__('POST'):
        email = request.form.get('email')
        if utils.get_user_by_email(email=email):
            return render_template('forgot_password.html', email=email)  # Gửi email qua template
        else:
            flash('Không có tài khoản nào đăng ký địa chỉ email này, hãy thử lại', 'error')
    return render_template('search_account.html')


@app.route('/forgotPassword', methods=['POST'])
def forgot_password():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    email = request.form.get('email')  # Lấy email từ trường ẩn
    password = request.form.get('password')
    confirm = request.form.get('confirm')
    if not password.__eq__(confirm):
        flash('Mật khẩu không khớp, vui lòng nhập lại', 'error')
        return render_template('forgot_password.html', email=email)
    else:
        user = utils.get_user_by_email(email=email)
        if user:
            user.password = hashlib.md5(password.strip().encode('utf-8')).hexdigest()
            db.session.add(user)
            db.session.commmit()
            flash('Mật khẩu đã được cập nhật thành công.', 'success')
            return redirect(url_for('login'))
    return render_template('forgot_password.html', email=email)


@app.route('/login_admin', methods=['POST'])
def login_admin():
    username = request.form.get('username')
    password = request.form.get('password')
    user = utils.check_login(username=username, password=password, role=NguoiDung.vaiTro)
    if user:
        login_user(user=user)

        return redirect('/admin')
    else:
        err_msg = 'username hoặc password ko chính xác'


@loginMNG.user_loader
def user_load(user_id):
    return utils.get_user_by_id(user_id=user_id)


@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))


@app.route('/profile')
@login_required
def profile():
    check = checkAuthenticated()
    if check:
        return check
    checkAuthenticated()
    return render_template('user_profile.html')


@app.route('/changeProfile', methods=['GET', 'POST'])
@login_required
def changeProfile():
    check = checkAuthenticated()
    if check:
        return check
    if request.method.__eq__('POST'):
        checkAuthenticated()
        name = request.form.get('name')
        phone = request.form.get('phone')
        address = request.form.get('address')
        email = request.form.get('email')
        gender = request.form.get('gender')
        birthday = request.form.get('birthday')
        avatar = request.files.get('avatar')
        if avatar:
            res = cloudinary.uploader.upload(avatar)
            avatar = res['secure_url']
            current_user.anhDaiDien = avatar

        current_user.hoVaTen = name
        current_user.sdt = phone
        current_user.diaChi = address
        current_user.email = email
        current_user.gioiTinh = 0 if gender == 'male' else 1
        current_user.ngaySinh = birthday

        db.session.commit()
        return redirect(url_for('profile'))

    return render_template('changeProfile.html')


@app.route("/changePassword", methods=['GET', 'POST'])
@login_required
def changePassword():
    check = checkAuthenticated()
    if check:
        return check
    if request.method.__eq__('POST'):

        old = request.form.get('old')
        new = request.form.get('new')
        confirm = request.form.get('confirm')

        if current_user.password.__eq__(str(hashlib.md5(old.strip().encode('utf-8')).hexdigest())):
            if new:
                if new.__eq__(confirm):
                    current_user.password = str(hashlib.md5(new.strip().encode('utf-8')).hexdigest())
                    db.session.commit()
                    flash('Đổi mật khẩu thành công.', 'success')
                else:
                    flash('Mật khẩu nhập lại không trùng khớp, mời nhập lại.', 'error')
            else:
                flash('Mật khẩu mới không hợp lệ, mời nhập lại.', 'error')
        else:
            flash('Mật khẩu cũ không đúng, mời nhập lại.', 'error')
    return render_template('changePassword.html')


@app.route("/products/<int:sach_id>", methods=['GET', 'POST'])
def productDetail(sach_id):
    book = utils.get_book_by_id(sach_id)
    a = utils.count_comment(sach_id)
    comments = utils.get_commment(int(request.args.get('page', 1)))
    if request.method == 'POST':
        data = request.get_json()  # Lấy dữ liệu JSON từ body của POST request
        print(data)
        quantity = data.get('quantity')  # Lấy giá trị số lượng, mặc định là 1 nếu không có
        if not current_user.is_authenticated:
            return redirect(url_for('login'))
        # Lấy giỏ hàng của người dùng
        gio_hang = GioHang.query.filter_by(id_NguoiDung=current_user.id).first()
        # #truy vấn sản phẩm theo id
        # product = Sach.query.filter_by(id=sach_id).first()
        if gio_hang is None:
            # Nếu không có giỏ hàng, tạo mới giỏ hàng cho người dùng
            gio_hang = GioHang(id_NguoiDung=current_user.id)
            db.session.add(gio_hang)
            db.session.commit()
        # Kiểm tra xem sản phẩm đã có trong giỏ hàng chưa
        chi_tiet = ChiTietGioHang.query.filter(
            ChiTietGioHang.id_GioHang == gio_hang.id,  # Đảm bảo `gio_hang` là một đối tượng
            ChiTietGioHang.id_Sach == sach_id).first()  # Đảm bảo `sach` là một đối tượng

        if chi_tiet:
            # Nếu sản phẩm đã có trong giỏ hàng, tăng số lượng lên
            chi_tiet.soLuong += int(quantity)
        else:
            # Nếu sản phẩm chưa có trong giỏ hàng, tạo mới chi tiết giỏ hàng
            chi_tiet = ChiTietGioHang(id_GioHang=gio_hang.id, id_Sach=sach_id, soLuong=int(quantity))
            db.session.add(chi_tiet)

        db.session.commit()

        # return redirect(url_for('add_to_cart', sach_id=sach_id))
        return jsonify({"message": "Sản phẩm đã được thêm vào giỏ hàng!"})
    return render_template('productDetail.html', sach=book, comments=comments,
                           pages=math.ceil(utils.count_comment(sach_id=sach_id) / app.config['COMMENT_SIZE']))


@app.route("/paymentNow", methods=['POST'])
def paymentNow():
    if not current_user.is_authenticated:
        return jsonify({"error": "Bạn cần đăng nhập để thanh toán."}), 401

    if request.method == 'POST':
        try:
            # Lấy dữ liệu từ body của yêu cầu
            data = request.get_json()
            print("Dữ liệu nhận được:", data)

            product = data.get('quantity')  # Giá trị mặc định nếu không có

            # Tạo đơn hàng
            don_hang = DonHang(id_NguoiDung=current_user.id)
            db.session.add(don_hang)
            db.session.commit()

            sach_id = product.get('id')
            quantity = product.get('quantity', 1)

            chi_tiet = ChiTietDonHang(
                id_DonHang=don_hang.id,
                id_Sach=sach_id,
                soLuong=int(quantity)
            )
            db.session.add(chi_tiet)

            db.session.commit()

            return jsonify({"message": "Thanh toán thành công!", "donHangId": don_hang.id})
        except Exception as e:
            print("Lỗi:", e)
            return jsonify({"error": str(e)}), 500
    return render_template('payment.html')


# note


@app.route("/import", methods=['GET', 'POST'])
@login_required
def import_book():
    check = checkImporter()
    if check:
        return check
    page = int(request.args.get('page', 1))
    size = 5
    start = (page - 1) * size
    end = start + size
    list_import = db.session.query(PhieuNhapSach.id, PhieuNhapSach.ngayNhapSach,
                                   NguoiDung.hoVaTen.label('hoVaTen')).join(NguoiDung,
                                                                            NguoiDung.id == PhieuNhapSach.id_NguoiDung).order_by(
        PhieuNhapSach.id.desc()).all()
    l = len(list_import)
    current_page = int(request.args.get('page', 1))
    list_import = list_import[start:end]
    list_import_detail = []
    selected_import = None

    if request.method == 'POST':
        data = request.json
        id_import = data.get('id_Import')
        if id_import:
            selected_import = utils.get_import_by_id(int(id_import))
            list_import_detail = db.session.query(
                ChiTietNhapSach.soLuong,
                Sach.ten.label("ten_sach"),
                Sach.tacGia.label("tac_gia"),
                TheLoai.ten.label("ten_the_loai")
            ).join(
                Sach, ChiTietNhapSach.id_Sach == Sach.id
            ).join(
                TheLoai, Sach.id_TheLoai == TheLoai.id
            ).filter(
                ChiTietNhapSach.id_PhieuNhapSach == id_import
            ).all()

            details = [{
                "ten_sach": ct.ten_sach,
                "ten_the_loai": ct.ten_the_loai,
                "tac_gia": ct.tac_gia,
                "soLuong": ct.soLuong
            } for ct in list_import_detail]

            return jsonify({
                "success": True,
                "details": details,
                "ngay_nhap": selected_import.ngayNhapSach.strftime('%d-%m-%Y') if selected_import else None
            })
        return jsonify({"success": False}), 400

    return render_template('import_home.html',
                           selected_import=selected_import,
                           list_import=list_import,
                           list_import_detail=list_import_detail,
                           page=math.ceil(l / 5),
                           current_page=current_page)


@app.route('/importCreate', methods=['GET', 'POST'])
@login_required
def create_import():
    check = checkImporter()
    if check:
        return check
    if request.method == 'GET':
        keyword = request.args.get('q')
        if keyword:
            books = Sach.query.filter(Sach.ten.like(f"%{keyword}%")).all()
            results = [
                {
                    'id': book.id,
                    'ten': book.ten,
                    'tacGia': book.tacGia,
                    'theLoai': book.TheLoai.ten,
                    'soLuongTonKho': book.soLuongTonKho
                }
                for book in books
            ]
            return jsonify(results)
        qd = utils.get_qd()
        return render_template('create_import.html', today=datetime.now().strftime('%Y-%m-%d'), quy_dinh=qd)

    elif request.method == 'POST':
        data = request.json
        ngay_nhap = data.get('ngayNhap')
        sach_list = data.get('sachList', [])
        phieu_nhap = PhieuNhapSach.query.filter(PhieuNhapSach.ngayNhapSach == ngay_nhap).first()
        if not phieu_nhap:
            db.session.add(PhieuNhapSach(ngayNhapSach=ngay_nhap, id_NguoiDung=current_user.id))
            db.session.commit()
            phieu_nhap = PhieuNhapSach.query.filter(PhieuNhapSach.ngayNhapSach == ngay_nhap).first()
        for sach in sach_list:
            chi_tiet = ChiTietNhapSach(
                soLuong=sach['soLuong'],
                id_Sach=sach['id_Sach'],
                id_PhieuNhapSach=phieu_nhap.id
            )
            db.session.add(chi_tiet)

            sach_obj = Sach.query.get(sach['id_Sach'])
            sach_obj.soLuongTonKho += sach['soLuong']

        db.session.commit()

        return jsonify({'success': True})


@app.route('/addBook', methods=['GET', 'POST'])
@login_required
def add_book():
    check = checkImporter()
    if check:
        return check
    err_msg = ''
    suc_msg = ''
    if request.method == 'POST':
        ten = request.form['ten']
        tacGia = request.form['tacGia']
        moTa = request.form.get('moTa', '')
        donGia = float(request.form['donGia'])
        id_TheLoai = int(request.form['id_TheLoai'])
        image = request.files['image']
        existing_book = Sach.query.filter_by(ten=ten, tacGia=tacGia).first()
        if existing_book:
            flash(f"Sách '{ten}' của tác giả '{tacGia}' đã tồn tại trong hệ thống.", 'danger')
            return redirect(url_for('add_book'))

        image_path = None
        if image:
            res = cloudinary.uploader.upload(image)
            image_path = res['secure_url']
        utils.add_book(ten=ten, tacGia=tacGia, donGia=donGia, moTa=moTa
                       , id_TheLoai=id_TheLoai, image=image_path)
        flash('Đã thêm sách thành công.', 'success')
        return redirect(url_for('add_book'))
    theLoaiList = TheLoai.query.all()
    return render_template('add_book.html', theLoaiList=theLoaiList)


@app.route('/bill', methods=['GET', 'POST'])
@login_required
def bill():
    check = checkEmployee()
    if check:
        return check
    page = int(request.args.get('page', 1))
    # list_bill = db.session.query(DonHang.id.label('id_don_hang'),
    #                              DonHang.ngayDatHang.label('ngay_dat_hang'),
    #                              DonHang.phuongThucThanhToan.label('pttt'),
    #                              NguoiDung.hoVaTen.label('ten_user'),
    #                              DonHang.trangThai.label('trang_thai'),
    #                              func.sum(Sach.donGia * ChiTietDonHang.soLuong).label('tong_gia')) \
    #     .select_from(DonHang).join(ChiTietDonHang, DonHang.id == ChiTietDonHang.id_DonHang) \
    #     .join(Sach, ChiTietDonHang.id_Sach == Sach.id) \
    #     .join(NguoiDung, DonHang.nguoiDung == NguoiDung.id) \
    #     .group_by(DonHang.id, DonHang.ngayDatHang, DonHang.phuongThucThanhToan, NguoiDung.hoVaTen,
    #               DonHang.trangThai).order_by(DonHang.id.desc()).all()
    list_bill = utils.get_bill_pagination(page)
    current_page = int(request.args.get('page', 1))
    selected_bill = None
    list_bill_detail = []
    if request.method == 'POST':
        data = request.json
        id_bill = data.get('id_Bill')
        if id_bill:
            selected_bill = utils.get_bill_by_id(int(id_bill))
            list_bill_detail = db.session.query(ChiTietDonHang.soLuong, Sach.ten.label("ten_sach"),
                                                Sach.tacGia.label("tac_gia"), TheLoai.ten.label("ten_the_loai")) \
                .join(Sach, ChiTietDonHang.id_Sach == Sach.id) \
                .join(TheLoai, Sach.id_TheLoai == TheLoai.id) \
                .filter(ChiTietDonHang.id_DonHang == id_bill) \
                .all()
        details = [{
            "ten_sach": ct.ten_sach,
            "ten_the_loai": ct.ten_the_loai,
            "tac_gia": ct.tac_gia,
            "soLuong": ct.soLuong
        } for ct in list_bill_detail]

        return jsonify({
            "success": True,
            "details": details,
            "ngay_dat_hang": selected_bill.ngayDatHang.strftime('%d-%m-%Y') if selected_bill else None
        })
    return render_template('bill_home.html', list_bill=list_bill,
                           selected_bill=selected_bill, list_bill_detail=list_bill_detail,
                           page=math.ceil(utils.count_bill_pagination() / app.config['LIST_SIZE']),
                           current_page=current_page)


# return render_template('products.html',
#                           page=math.ceil(utils.count_book(category_id=cate_id, kw=kw) / app.config['LIST_SIZE']),
#                           lb=sach)


@app.route('/billCreate', methods=['GET', 'POST'])
@login_required
def create_bill():
    check = checkEmployee()
    if check:
        return check
    if request.method == 'GET':
        id = request.args.get('id')
        if id:
            books = Sach.query.filter(Sach.id == int(id)).all()
            results = [
                {
                    'id': book.id,
                    'ten': book.ten,
                    'theLoai': book.TheLoai.ten,
                    'donGia': book.donGia
                }
                for book in books
            ]
            return jsonify(results)

        query = request.args.get('q')
        if query:
            users = NguoiDung.query.filter(NguoiDung.hoVaTen.ilike(f"%{query}%")).all()
            user_results = [
                {
                    'id': user.id,
                    'ten': user.hoVaTen,
                    'email': user.email
                }
                for user in users
            ]
            return jsonify(user_results)
        return render_template('create_bill.html', today=datetime.today().date())
    elif request.method == 'POST':
        data = request.json
        ngay_lap_hoa_don = data.get('ngayDatHang')
        sach_list = data.get('sachList', [])
        id = data.get('id')
        don_hang = DonHang(
            ngayDatHang=ngay_lap_hoa_don,
            ngayThanhToan=ngay_lap_hoa_don,
            trangThai=TrangThai.DA_NHAN_HANG,
            phuongThucThanhToan=PhuongThucThanhToan.TRUC_TIEP,
            id_NguoiDung=current_user.id if not id else utils.get_user_by_id(user_id=id).id
        )
        db.session.add(don_hang)
        db.session.commit()

        hoa_don = DonHangOffline(id_NhanVien=current_user.id, id_DonHang=don_hang.id)
        db.session.add(hoa_don)
        db.session.commit()


        for p in sach_list:
            chi_tiet = ChiTietDonHang(
                id_Sach=p['id_Sach'],
                id_DonHang=don_hang.id,
                soLuong=p['soLuong']
            )
            db.session.add(chi_tiet)

            sach_obj = utils.get_book_by_id(p['id_Sach'])
            sach_obj.soLuongTonKho -= p['soLuong']

        db.session.commit()

        return jsonify({"success": True})


@app.route('/billExport', methods=['POST'])
@login_required
def export_bill():
    data = request.json
    ngay_dat_hang = data.get('ngayDatHang')
    id = utils.count_bill() + 1
    id_kh = data.get('id_kh')
    sach_list = data.get('sachList', [])
    ten_kh = utils.get_user_by_id(user_id=id_kh) if id_kh else None
    tongTien, tongSoLuong = 0, 0
    document = Document()

    document.add_heading('Book Store', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('Hóa đơn bán sách', style='Heading 2').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('-' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f'Mã hóa đơn: {id}').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if id_kh:
        document.add_paragraph(f'Tên khách hàng: {ten_kh}')
    document.add_paragraph(f'Ngày lập hóa đơn: {ngay_dat_hang}')

    table = document.add_table(rows=1, cols=5, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tên sách'
    hdr_cells[1].text = 'Số lượng'
    hdr_cells[2].text = 'Đơn giá'
    hdr_cells[3].text = 'Thành tiền'

    # Điền dữ liệu vào bảng
    for sach in sach_list:
        book = utils.get_book_by_id(sach['id_Sach'])
        row_cells = table.add_row().cells
        row_cells[0].text = book.ten
        row_cells[1].text = str(sach['soLuong'])
        row_cells[2].text = f"{book.donGia:,} VND"
        thanhTien = book.donGia * sach['soLuong']
        row_cells[3].text = f"{thanhTien:,} VND"
        tongTien += thanhTien
        tongSoLuong += sach['soLuong']

    # Thêm tổng cộng bên dưới bảng
    document.add_paragraph('-' * 50)
    document.add_paragraph(f'Số lượng: {tongSoLuong} sản phẩm')
    document.add_paragraph(f'Thành tiền: {tongTien:,} VND')
    document.add_paragraph('-' * 50)

    # Thêm thông tin nhân viên và lời cảm ơn
    document.add_paragraph('Họ tên nhân viên: Tấn Phạm')
    document.add_paragraph('Thank you & see you again!')
    document.add_paragraph('Xã Nhơn Đức, huyện Nhà Bè, TP.HCM')

    # table = document.add_table(rows=1, cols=5)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Mã sách'
    # hdr_cells[1].text = 'Tên sách'
    # hdr_cells[2].text = 'Thể loại'
    # hdr_cells[3].text = 'Đơn giá'
    # hdr_cells[4].text = 'Số lượng'
    #
    # for sach in sach_list:
    #     book = utils.get_book_by_id(sach['id_Sach'])
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(book.id)
    #     row_cells[1].text = book.ten
    #     row_cells[2].text = book.TheLoai.ten
    #     row_cells[3].text = f"{book.donGia:,} VND"
    #     row_cells[4].text = str(sach['soLuong'])
    #     tongTien += book.donGia * sach['soLuong']
    #
    # for row in table.rows:
    #     for cell in row.cells:
    #         # Đảm bảo tất cả các ô có viền
    #         cell._element.get_or_add_tcPr().append(
    #             parse_xml(
    #                 f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:space="0" w:color="000000" w:sz="4"/><w:left w:val="single" w:space="0" w:color="000000" w:sz="4"/><w:bottom w:val="single" w:space="0" w:color="000000" w:sz="4"/><w:right w:val="single" w:space="0" w:color="000000" w:sz="4"/></w:tcBorders>'
    #             )
    #         )

    os.makedirs('exports', exist_ok=True)

    document.add_paragraph(f'Thành tiền: {tongTien}')

    # Lưu file vào server
    file_name = f'HoaDon_{id}.docx'
    file_path = os.path.join('exports', file_name)
    document.save(file_path)

    # Trả về đường dẫn tải file
    download_url = url_for('download_file', filename=file_name, _external=True)
    return jsonify({
        'success': True,
        'download_url': download_url
    })


@app.route('/download/<filename>')
def download_file(filename):
    return send_file(f'exports/{filename}', as_attachment=True)


@app.route('/api/comments', methods=['POST'])
@login_required
def add_comment():
    data = request.json
    content = data.get('content')  # Sửa ở đây
    sach_id = data.get('sach_id')  # Sửa ở đây

    try:
        c = utils.add_comment(noiDung=content, sach_id=sach_id)
    except:
        return {'status': 404, 'err_msg': 'Đã có lỗi xảy ra'}

    return {'status': 201, 'comment': {
        'id': c.id,
        'content': c.noiDung,
        'created_date': c.thoiGian,
        'user': {
            'name': current_user.hoVaTen,
            'avatar': current_user.anhDaiDien
        }
    }}


def checkAuthenticated():
    if not current_user.is_authenticated:
        return redirect(url_for('index'))


def checkEmployee():
    if not current_user.is_authenticated or not current_user.is_employee():
        return redirect(url_for('index'))


def checkImporter():
    if not current_user.is_authenticated or not current_user.is_importer():
        return redirect(url_for('index'))


@app.route('/cart', methods=['GET', 'POST'])
def cart():
    if not current_user.is_authenticated:
        return redirect('/login')
        # Giả sử bạn lấy đơn hàng từ database dựa trên user_id
    current_user_id = current_user.id  # ID người dùng hiện tại (nếu có đăng nhập)

    # lấy giỏ hàng của người dùng đang đăng nhập hiện tại ( mỗi người chỉ có 1 giỏ hàng)
    gio_hang = GioHang.query.filter_by(id_NguoiDung=current_user_id).first()
    if not gio_hang:
        return render_template('cart.html', carts=[])  # nếu giỏ hàng rỗng thì trả về rỗng
    # lấy danh sách chi tiết giỏ hàng của người dùng
    chi_tiet_GH = ChiTietGioHang.query.filter_by(id_GioHang=gio_hang.id).all()

    danh_sach_sach = []
    for chi_tiet in chi_tiet_GH:
        danh_sach_sach.append({
            'id': chi_tiet.Sach.id,
            'image': chi_tiet.Sach.image,
            'ten': chi_tiet.Sach.ten,
            'tacGia': chi_tiet.Sach.tacGia,
            'soLuong': chi_tiet.soLuong,
            'donGia': chi_tiet.Sach.donGia,
            'tongTien': chi_tiet.soLuong * chi_tiet.Sach.donGia,  # Tính tổng tiền
        })
    # Truyền dữ liệu vào template giỏ hàng
    return render_template('cart.html', carts=danh_sach_sach)


# @app.route('/api/pay')
# @login_required
# def pay():
#     # ghi nhận đơn hàng
#     key = app.config['CART_KEY']  # lấy 'cart' ra
#     cart = session.get(key)
#
#     try:  # check có bị lỗi ko, thường bị lỗi ràng buộc về CSDL
#         dao.save_order(cart)
#     except Exception as ex:  # lỗi thì vào đây
#         print(str(ex))
#         return jsonify({
#             "status": 500,
#             "message": "Lỗi hệ thống",
#         })
#     else:  # chạy lệnh success
#         # del session(key) # xóa cart trong session, vì đã thanh toán xong
#         pass
#
#     return jsonify({
#         "status": 200,
#         "message": "Hoàn tất thanh toán",
#     })


@app.route('/products/<int:sach_id>', methods=['POST'])
@login_required
def add_to_cart(sach_id):
    return render_template('productDetail.html')  # Chuyển hướng về trang giỏ hàng


@app.route('/payment', methods=['GET', 'POST'])
@login_required
def payment():
    # Kiểm tra Content-Type
    if request.method == 'POST':
        products = []
        data = request.get_json()  # Lấy dữ liệu JSON từ body của POST request
        print("Data received:", data)

        # Lấy giá trị của `products`
        raw_products = data.get('products', '[]')  # Lấy giá trị mặc định là chuỗi '[]'

        try:
            # Parse chuỗi JSON thành danh sách Python
            products = json.loads(raw_products)
            print("Parsed products:", products)
        except json.JSONDecodeError as e:
            print("Error decoding JSON:", e)
        # lây pttt
        paymentMethod = products[0].get('paymentMethod', 'TRUC_TIEP')
        if paymentMethod == 'TRUC_TIEP':
            paymentMethod = PhuongThucThanhToan.TRUC_TIEP
        else:
            paymentMethod = PhuongThucThanhToan.TRUC_TUYEN
        don_hang = DonHang(id_NguoiDung=current_user.id, phuongThucThanhToan=paymentMethod)
        db.session.add(don_hang)
        db.session.commit()
        for pro in products:
            if isinstance(pro, dict):  # Đảm bảo mỗi phần tử là dictionary
                sach_id = pro.get('id')
                quantity = pro.get('quantity')
                if pro.get('paymentMethod') == 'Thanh toán trực tiếp':
                    paymentMethod = 'TRUC_TIEP'
                if pro.get('paymentMethod') == 'Thanh toán Momo':
                    paymentMethod = 'TRUC_TUYEN'
                print(f"sach_id: {sach_id}, quantity: {quantity}")
                chi_tiet = ChiTietDonHang(id_DonHang=don_hang.id, id_Sach=sach_id, soLuong=int(quantity))
                db.session.add(chi_tiet)
                db.session.commit()
            else:
                print("Invalid product data:", pro)
        print(don_hang.id)
        return jsonify({"message": "Thanh toán thành công!", "donHangId": don_hang.id})
    return render_template('payment.html')
    # don_hang = DonHang.query.filter_by(nguoiDung=current_user.id).order_by(
    #     DonHang.id.desc()).first()  # Lấy đơn hàng mới nhất của người dùng
    # chi_tiet_DH = ChiTietDonHang.query.filter_by(id_DonHang=don_hang.id).all()
    #
    # danh_sach_sach = []
    # for chi_tiet in chi_tiet_DH:
    #     danh_sach_sach.append({
    #         'id': chi_tiet.Sach.id,
    #         'image': chi_tiet.Sach.image,
    #         'ten': chi_tiet.Sach.ten,
    #         'tacGia': chi_tiet.Sach.tacGia,
    #         'soLuong': chi_tiet.soLuong,
    #         'donGia': chi_tiet.Sach.donGia,
    #         'tongTien': chi_tiet.soLuong * chi_tiet.Sach.donGia,  # Tính tổng tiền
    #     })
    # print(danh_sach_sach)


@app.route('/status')
def status():
    return render_template('status.html', orders=utils.get_order())


@app.route('/api/pay', methods=['POST'])
def api_pay():
    try:
        data = request.json
        amount = int(data.get('amount', 100000))
        chi_tiet_don_hang = data.get('chi_tiet_don_hang')
        don_hang = DonHang(
            trangThai=TrangThai.DANG_CHO_THANH_TOAN,
            phuongThucThanhToan=PhuongThucThanhToan.TRUC_TUYEN,
            id_NguoiDung=current_user.id
        )
        db.session.add(don_hang)
        db.session.commit()
        # vnp = Vnpay()
        # vnp.requestData = {
        #     'vnp_Version': '2.1.0',
        #     'vnp_Command': 'pay',
        #     'vnp_TmnCode': ,  # Kiểm tra mã TMN code
        #     'vnp_Amount': amount * 100,  # Nhân 100 để đưa về VNĐ
        #     'vnp_CurrCode': 'VND',
        #     'vnp_TxnRef': str(don_hang.id),  # Mã đơn hàng duy nhất
        #     'vnp_OrderInfo': str(don_hang.id),
        #     'vnp_OrderType': 'billpayment',
        #     'vnp_Locale': 'vn',
        #     'vnp_BankCode': 'NCB',
        #     'vnp_ReturnUrl': app.config["VNPAY_RETURN_URL"]  # URL trả về sau thanh toán
        # }
        # print( vnp.requestData)
        ipaddr = request.remote_addr
        # Build URL Payment
        vnp = Vnpay()
        vnp.requestData['vnp_Version'] = '2.1.0'
        vnp.requestData['vnp_Command'] = 'pay'
        vnp.requestData['vnp_TmnCode'] = app.config["VNPAY_TMN_CODE"]
        vnp.requestData['vnp_Amount'] = amount * 100
        vnp.requestData['vnp_CurrCode'] = 'VND'
        vnp.requestData['vnp_TxnRef'] = str(don_hang.id) + "_" + datetime.now().__str__()
        vnp.requestData['vnp_OrderInfo'] = str(don_hang.id)
        vnp.requestData['vnp_OrderType'] = str(don_hang.id)
        # Check language, default: vn
        vnp.requestData['vnp_Locale'] = 'vn'
        vnp.requestData['vnp_BankCode'] = 'NCB'

        vnp.requestData['vnp_CreateDate'] = datetime.now().strftime('%Y%m%d%H%M%S')  # 20150410063022
        vnp.requestData['vnp_IpAddr'] = ipaddr
        vnp.requestData['vnp_ReturnUrl'] = app.config["VNPAY_RETURN_URL"]

        for p in chi_tiet_don_hang:
            chi_tiet = ChiTietDonHang(
                id_Sach=p['product_id'],
                id_DonHang=don_hang.id,
                soLuong=p['quantity'],
            )
            ChiTietGioHang.query.filter_by(
                id_Sach=p['product_id'],
                id_GioHang=GioHang.query.filter(GioHang.id_NguoiDung==current_user.id).first().id
            ).delete()
            db.session.add(chi_tiet)
        db.session.commit()
        print("Request Data:", vnp.requestData)
        payment_url = vnp.get_payment_url(
            'https://sandbox.vnpayment.vn/paymentv2/vpcpay.html',
            app.config["VNPAY_HASH_SECRET_KEY"]
        )

        # Log URL thanh toán
        print("Payment URL:", payment_url)

        return jsonify({'status': 200, 'payment_url': payment_url})
    except Exception as e:
        return jsonify({'status': 500, 'message': str(e)})



@app.route("/payment_return", methods=["GET"])
def payment_return():
    if request.args:
        vnp = Vnpay()
        vnp.responseData = request.args.to_dict()
        order_id = request.args.get('vnp_TxnRef')
        amount = int(request.args.get('vnp_Amount')) / 100
        order_desc = request.args.get('vnp_OrderInfo')
        vnp_BankTranNo = request.args.get("vnp_BankTranNo")
        vnp_TransactionNo = request.args.get('vnp_TransactionNo')
        vnp_ResponseCode = request.args.get('vnp_ResponseCode')
        print(vnp_ResponseCode)
        vnp_PayDate = request.args.get('vnp_PayDate')
        vnp_BankCode = request.args.get('vnp_BankCode')
        vnp_CardType = request.args.get('vnp_CardType')
        vnp_SecureHash = request.args.get('vnp_SecureHash')
        if vnp.validate_response(app.config["VNPAY_HASH_SECRET_KEY"]):
            if vnp_ResponseCode == "00":
                dh=DonHang.query.get(order_id)
                dh.trangThai=TrangThai.DANG_GIAO_HANG
                db.session.commit()
                return render_template("vnpay/payment_return.html", title="Kết quả giao dịch",
                                       result="Thành công", order_id=order_id,
                                       amount=amount,
                                       order_desc=order_desc,
                                       vnp_TransactionNo=vnp_TransactionNo,
                                       vnp_ResponseCode=vnp_ResponseCode)
            else:
                return render_template("vnpay/payment_return.html", title="Kết quả giao dịch",
                                       result="Lỗi", order_id=order_id,
                                       amount=amount,
                                       order_desc=order_desc,
                                       vnp_TransactionNo=vnp_TransactionNo,
                                       vnp_ResponseCode=vnp_ResponseCode)
        else:
            return render_template("vnpay/payment_return.html",
                                   title="Kết quả giao dịch", result="Lỗi", order_id=order_id, amount=amount,
                                   order_desc=order_desc, vnp_TransactionNo=vnp_TransactionNo,
                                   vnp_ResponseCode=vnp_ResponseCode, msg="Sai checksum")
    else:
        return render_template("vnpay/payment_return.html", title="Kết quả thanh toán", result="")


def cap_nhat_trang_thai_don_hang():
    try:
        # Thời điểm hiện tại
        now = datetime.now()

        # Tìm các đơn hàng cần cập nhật
        don_hang_can_cap_nhat = DonHang.query.filter(
            DonHang.trangThai == TrangThai.DANG_CHO_THANH_TOAN,
            DonHang.ngayDatHang <= now - timedelta(hours=48)
        ).all()

        # Cập nhật trạng thái cho từng đơn hàng
        for don_hang in don_hang_can_cap_nhat:
            don_hang.trangThai = TrangThai.DA_HUY
            print(f"Cập nhật đơn hàng ID {don_hang.id}: Chuyển sang trạng thái 'Hủy'")

        # Lưu thay đổi vào cơ sở dữ liệu
        db.session.commit()
        print(f"Cập nhật trạng thái cho {len(don_hang_can_cap_nhat)} đơn hàng.")
    except Exception as e:
        db.session.rollback()  # Hoàn tác nếu có lỗi
        print(f"Lỗi khi cập nhật trạng thái đơn hàng: {e}")


def run_scheduler():
    with app.app_context():
        cap_nhat_trang_thai_don_hang()

# from apscheduler.schedulers.background import BackgroundScheduler
#
# def start_scheduler():
#     scheduler = BackgroundScheduler()
#     scheduler.add_job(run_scheduler, 'interval', hours=1)
#     scheduler.start()
#
#     try:
#         while True:
#             pass
#     except (KeyboardInterrupt, SystemExit):
#         scheduler.shutdown()

if __name__ == '__main__':
    with app.app_context():
        app.run(debug=True)
